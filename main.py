import docx
import re

doc = docx.Document('D:\第1章考題參考-80題.docx')
# print('段落數量： ', len(doc.paragraphs))
paraStr = ""
splitList = []
    
for para in doc.paragraphs:
    paraStr += para.text
paraStr = re.sub('\n', '', paraStr)
paraStr = re.sub(' ', '', paraStr)
paraStr = re.sub('。', '', paraStr)
paraStr = re.sub('？','?',paraStr)
paraStr = re.sub('：', '?', paraStr)
paraStr = re.sub(':', '?', paraStr)
# print(paraStr)
topic_pattern = r'\d{1,2}\..*?\?'
answer_pattern = r'答案\?([a-z])..*?答案解釋'

topic = re.findall(topic_pattern, paraStr)

answer = re.findall(answer_pattern, paraStr)
for a in answer:
    print(a)
    print()
print(len(answer))